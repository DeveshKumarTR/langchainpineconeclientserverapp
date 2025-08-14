from flask import Blueprint, request, jsonify, current_app
from werkzeug.utils import secure_filename
import os
import uuid
from datetime import datetime

from server.models.vector_store import VectorStoreManager
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_core.documents import Document

documents_bp = Blueprint('documents', __name__)

def allowed_file(filename):
    """Check if file extension is allowed"""
    return '.' in filename and \
           filename.rsplit('.', 1)[1].lower() in current_app.config['ALLOWED_EXTENSIONS']

def load_document(file_path, filename):
    """Load document based on file type - simplified version for Windows compatibility"""
    file_ext = filename.rsplit('.', 1)[1].lower()
    
    try:
        if file_ext == 'txt':
            # Simple text file loader
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            return [Document(page_content=content, metadata={"source": file_path})]
        
        elif file_ext == 'pdf':
            try:
                import PyPDF2
                documents = []
                with open(file_path, 'rb') as f:
                    pdf_reader = PyPDF2.PdfReader(f)
                    for page_num, page in enumerate(pdf_reader.pages):
                        text = page.extract_text()
                        if text.strip():
                            documents.append(Document(
                                page_content=text,
                                metadata={"source": file_path, "page": page_num}
                            ))
                return documents
            except ImportError:
                raise ValueError("PyPDF2 is required for PDF processing")
        
        elif file_ext == 'docx':
            try:
                import docx
                doc = docx.Document(file_path)
                content = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
                return [Document(page_content=content, metadata={"source": file_path})]
            except ImportError:
                raise ValueError("python-docx is required for DOCX processing")
        
        elif file_ext == 'xlsx':
            try:
                import openpyxl
                workbook = openpyxl.load_workbook(file_path)
                content = ""
                for sheet_name in workbook.sheetnames:
                    sheet = workbook[sheet_name]
                    content += f"Sheet: {sheet_name}\n"
                    for row in sheet.iter_rows(values_only=True):
                        row_text = '\t'.join([str(cell) if cell is not None else '' for cell in row])
                        content += row_text + '\n'
                    content += '\n'
                return [Document(page_content=content, metadata={"source": file_path})]
            except ImportError:
                raise ValueError("openpyxl is required for XLSX processing")
        
        else:
            raise ValueError(f"Unsupported file type: {file_ext}")
            
    except Exception as e:
        raise ValueError(f"Error loading {file_ext} file: {str(e)}")

@documents_bp.route('/documents', methods=['POST'])
def upload_document():
    """Upload and process a document"""
    try:
        # Check if file is present
        if 'file' not in request.files:
            return jsonify({'error': 'No file provided'}), 400
        
        file = request.files['file']
        if file.filename == '':
            return jsonify({'error': 'No file selected'}), 400
        
        if not allowed_file(file.filename):
            return jsonify({'error': 'File type not allowed'}), 400
        
        # Generate unique filename and save
        filename = secure_filename(file.filename)
        doc_id = str(uuid.uuid4())
        temp_path = f"/tmp/{doc_id}_{filename}"
        file.save(temp_path)
        
        try:
            # Load and split document
            documents = load_document(temp_path, filename)
            
            text_splitter = RecursiveCharacterTextSplitter(
                chunk_size=current_app.config['CHUNK_SIZE'],
                chunk_overlap=current_app.config['CHUNK_OVERLAP']
            )
            texts = text_splitter.split_documents(documents)
            
            # Add metadata
            for text in texts:
                text.metadata.update({
                    'doc_id': doc_id,
                    'filename': filename,
                    'upload_time': datetime.utcnow().isoformat()
                })
            
            # Store in vector database
            vector_store = VectorStoreManager()
            vector_store.add_documents(texts)
            
            return jsonify({
                'success': True,
                'doc_id': doc_id,
                'filename': filename,
                'chunks_created': len(texts),
                'message': 'Document processed and stored successfully'
            }), 201
            
        finally:
            # Clean up temporary file
            if os.path.exists(temp_path):
                os.remove(temp_path)
            
    except Exception as e:
        current_app.logger.error(f"Error processing document: {str(e)}")
        return jsonify({'error': f'Failed to process document: {str(e)}'}), 500

@documents_bp.route('/documents', methods=['GET'])
def list_documents():
    """List all processed documents"""
    try:
        vector_store = VectorStoreManager()
        documents = vector_store.list_documents()
        
        return jsonify({
            'success': True,
            'documents': documents
        }), 200
        
    except Exception as e:
        current_app.logger.error(f"Error listing documents: {str(e)}")
        return jsonify({'error': f'Failed to list documents: {str(e)}'}), 500

@documents_bp.route('/documents/<doc_id>', methods=['DELETE'])
def delete_document(doc_id):
    """Delete a document and all its chunks"""
    try:
        vector_store = VectorStoreManager()
        success = vector_store.delete_document(doc_id)
        
        if success:
            return jsonify({
                'success': True,
                'message': f'Document {doc_id} deleted successfully'
            }), 200
        else:
            return jsonify({'error': 'Document not found'}), 404
            
    except Exception as e:
        current_app.logger.error(f"Error deleting document: {str(e)}")
        return jsonify({'error': f'Failed to delete document: {str(e)}'}), 500
